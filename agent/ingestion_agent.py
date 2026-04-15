import json
import re
import hashlib
from pathlib import Path
from typing import Any, Dict, Optional

import pdfplumber
import docx
import chardet
import ollama as ollama_client
import structlog

from agents.base_agent import BaseAgent
from shared.models import ParsedCV, FileFormat
from shared.utils import truncate_text, get_file_extension
from config.settings import settings

logger = structlog.get_logger(__name__)


class IngestionAgent(BaseAgent):
    """
    Responsible for:
    1. Detecting file format
    2. Extracting raw text from PDF/DOCX/TXT
    3. Calling Ollama LLM to extract structured fields
    4. Returning a ParsedCV object with confidence score
    5. Logging ALL failures without stopping bulk batches
    """

    def __init__(self, db):
        super().__init__("ingestion", db)
        self.ollama = ollama_client.AsyncClient(host=settings.ollama_base_url)

    async def execute(self, payload: Dict[str, Any], correlation_id: str) -> Dict[str, Any]:
        """
        payload expects:
          - file_path: str     — path to the uploaded file
          - filename:  str     — original filename
        """
        file_path = payload["file_path"]
        filename  = payload["filename"]

        self.log.info("ingestion_start", filename=filename, path=file_path)

        # ---- Step 1: Detect format and extract raw text ----
        ext = get_file_extension(filename)
        raw_text, parse_warnings = await self._extract_text(file_path, ext)

        if not raw_text or len(raw_text.strip()) < 50:
            raise ValueError(f"Could not extract meaningful text from {filename}")

        # ---- Step 2: Compute file hash (deduplication) ----
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        file_hash = hashlib.sha256(file_bytes).hexdigest()

        # ---- Step 3: LLM extraction ----
        truncated_text = truncate_text(raw_text, max_chars=6000)
        parsed_cv = await self._llm_extract(truncated_text, parse_warnings)

        self.log.info(
            "ingestion_complete",
            filename=filename,
            confidence=parsed_cv.parse_confidence,
            skills_found=len(parsed_cv.technical_skills),
        )

        return {
            "parsed_cv":      parsed_cv.model_dump(),
            "raw_text":       raw_text,
            "file_hash":      file_hash,
            "file_format":    ext,
            "parse_warnings": parse_warnings,
        }

    async def _extract_text(self, file_path: str, ext: str) -> tuple[str, list]:
        """Dispatch to the right extractor based on file type."""
        warnings = []

        if ext == "pdf":
            return await self._extract_pdf(file_path, warnings), warnings
        elif ext == "docx":
            return self._extract_docx(file_path, warnings), warnings
        elif ext == "txt":
            return self._extract_txt(file_path, warnings), warnings
        else:
            warnings.append(f"Unsupported format: {ext}")
            return "", warnings

    async def _extract_pdf(self, path: str, warnings: list) -> str:
        """Extract text from PDF using pdfplumber (handles tables & columns)."""
        text_parts = []
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        warnings.append(f"Page {i+1}: No extractable text (possibly scanned image)")
        except Exception as e:
            warnings.append(f"PDF extraction error: {str(e)}")

        return "\n".join(text_parts)

    def _extract_docx(self, path: str, warnings: list) -> str:
        """Extract text from DOCX preserving paragraph structure."""
        try:
            doc = docx.Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs)
        except Exception as e:
            warnings.append(f"DOCX extraction error: {str(e)}")
            return ""

    def _extract_txt(self, path: str, warnings: list) -> str:
        """Extract text from TXT with encoding detection."""
        try:
            with open(path, "rb") as f:
                raw_bytes = f.read()
            detected = chardet.detect(raw_bytes)
            encoding = detected.get("encoding", "utf-8") or "utf-8"
            return raw_bytes.decode(encoding, errors="replace")
        except Exception as e:
            warnings.append(f"TXT extraction error: {str(e)}")
            return ""

    async def _llm_extract(self, text: str, existing_warnings: list) -> ParsedCV:
        """
        Use Ollama LLM to extract structured fields from CV text.
        Returns a ParsedCV with confidence score.

        GUARDRAILS enforced in prompt:
        - Do NOT extract age, gender, photo description, or nationality
        - Evaluate only skills and experience
        """
        extraction_prompt = f"""
You are an expert HR data extraction assistant. Extract structured information from the CV text below.

IMPORTANT GUARDRAILS:
- Do NOT include age, date of birth, gender, nationality, religion, or any protected attributes
- Extract ONLY professional and skills-based information
- If a field is not present, use null

Return a valid JSON object with EXACTLY this structure:
{{
  "full_name": "string or null",
  "email": "string or null",
  "phone": "string or null",
  "linkedin_url": "string or null",
  "location": "city and country only, e.g. Dubai, UAE",
  "technical_skills": ["list of technical skills"],
  "soft_skills": ["list of soft skills like leadership, communication"],
  "domain_expertise": ["list of domains, e.g. cloud computing, fintech, supply chain"],
  "certifications": ["list of certifications"],
  "languages": ["spoken languages only, not programming languages"],
  "experience": [
    {{
      "company": "company name",
      "role": "job title",
      "start_date": "YYYY-MM or YYYY",
      "end_date": "YYYY-MM or YYYY or Present",
      "duration_months": integer or null,
      "description": "brief role summary",
      "technologies": ["tech used in this role"]
    }}
  ],
  "education": [
    {{
      "institution": "university or school name",
      "degree": "degree type e.g. BSc, MBA",
      "field": "field of study",
      "graduation_year": integer or null,
      "grade": "GPA or grade if mentioned"
    }}
  ],
  "total_years_exp": float or null,
  "cv_summary": "One paragraph professional summary of this candidate",
  "parse_confidence": float between 0.0 and 1.0 indicating how complete/clear the CV was
}}

Return ONLY the JSON object. No explanation. No markdown. No additional text.

CV TEXT:
{text}
"""
        try:
            response = await self.ollama.chat(
                model=settings.ollama_llm_model,
                messages=[{"role": "user", "content": extraction_prompt}],
                options={"temperature": 0.1, "num_predict": 2000},
            )

            raw_json = response["message"]["content"].strip()

            # Clean up common LLM formatting issues
            raw_json = re.sub(r"```json\s*", "", raw_json)
            raw_json = re.sub(r"```\s*", "", raw_json)

            data = json.loads(raw_json)
            return ParsedCV(**data)

        except json.JSONDecodeError as e:
            self.log.warning("llm_json_parse_error", error=str(e))
            existing_warnings.append(f"LLM returned malformed JSON: {str(e)}")
            return ParsedCV(parse_confidence=0.2, parse_warnings=existing_warnings)

        except Exception as e:
            self.log.error("llm_extraction_failed", error=str(e))
            existing_warnings.append(f"LLM extraction failed: {str(e)}")
            return ParsedCV(parse_confidence=0.1, parse_warnings=existing_warnings)
